"""
Document preprocessing for Honora EMS.
Supports PDF (.pdf) and Word (.docx) files.

Extraction strategy:
  PDF  – PyMuPDF find_tables() for tables; body text extracted with table
          regions excluded to avoid double-counting.
  DOCX – python-docx: paragraphs for body text, document tables for table data.
         Both use the same serialization and cleaning pipeline.

Public API
----------
extract_and_clean(source, filename)  -> str
extract_full(source, filename)       -> ExtractionResult
    .text        str                  # body text only (cleaned)
    .tables      list[TableData]      # structured table objects
    .combined    str                  # text + serialized tables (for embedding)
chunk_text(text, ...)                -> list[str]
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import fitz                      # PyMuPDF ≥ 1.23
from docx import Document        # python-docx

# ── Legal term normalisations ──────────────────────────────────────────────
LEGAL_TERM_MAP: dict[str, str] = {
    r"\bplaintiff\b": "plaintiff",
    r"\bdefendant\b": "defendant",
    r"\baffidavit\b": "affidavit",
    r"\bexhibit\b": "exhibit",
    r"\bdeposition\b": "deposition",
    r"\bsubpoena\b": "subpoena",
    r"\bjurisdiction\b": "jurisdiction",
    r"\bindictment\b": "indictment",
    r"\bprosecution\b": "prosecution",
    r"\bwitness\b": "witness",
}

# ── Data structures ────────────────────────────────────────────────────────

@dataclass
class TableData:
    page: int                          # 0-based page / table-ordinal in docx
    table_index: int
    headers: list[str]
    rows: list[list[str]]
    serialized: str = field(default="", repr=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "page": self.page,
            "table_index": self.table_index,
            "headers": self.headers,
            "rows": self.rows,
            "serialized": self.serialized,
        }


@dataclass
class ExtractionResult:
    text: str
    tables: list[TableData]
    combined: str


# ── Shared helpers ─────────────────────────────────────────────────────────

def _cell(val: Any) -> str:
    if val is None:
        return ""
    return re.sub(r"\s+", " ", str(val)).strip()


def _serialize_table(
    headers: list[str],
    rows: list[list[str]],
    page: int,
    t_idx: int,
) -> str:
    col_labels = headers if any(h for h in headers) else [
        f"Column {i+1}" for i in range(len(headers) or (len(rows[0]) if rows else 0))
    ]
    lines: list[str] = [f"[Table {t_idx + 1}, page {page + 1}]"]
    if any(h for h in col_labels):
        lines.append("Headers: " + " | ".join(col_labels) + ".")
    for r_idx, row in enumerate(rows):
        if not any(v for v in row):
            continue
        pairs = ", ".join(
            f"{label}={val}"
            for label, val in zip(col_labels, row)
            if val
        )
        lines.append(f"Row {r_idx + 1}: {pairs}.")
    return "\n".join(lines)


def clean_text(text: str) -> str:
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    for pattern, replacement in LEGAL_TERM_MAP.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text.strip()


# ── PDF extraction ─────────────────────────────────────────────────────────

def _table_bboxes(page: fitz.Page) -> list[fitz.Rect]:
    try:
        tabs = page.find_tables()
        return [t.bbox for t in tabs.tables] if tabs.tables else []
    except Exception:
        return []


def _rect_overlaps(r: fitz.Rect, bboxes: list[fitz.Rect]) -> bool:
    return any(not r.intersect(b).is_empty for b in bboxes)


def _pdf_body_text(page: fitz.Page, table_bboxes: list[fitz.Rect]) -> str:
    if not table_bboxes:
        return page.get_text("text")
    blocks = page.get_text("blocks")
    parts: list[str] = []
    for blk in blocks:
        if blk[6] != 0:
            continue
        if not _rect_overlaps(fitz.Rect(blk[:4]), table_bboxes):
            parts.append(blk[4])
    return "\n".join(parts)


def _pdf_tables(page: fitz.Page, page_idx: int) -> list[TableData]:
    try:
        tab_finder = page.find_tables()
        raw_tables = tab_finder.tables if tab_finder.tables else []
    except Exception:
        return []
    result: list[TableData] = []
    for t_idx, tab in enumerate(raw_tables):
        raw = tab.extract()
        if not raw:
            continue
        grid = [[_cell(c) for c in row] for row in raw]
        headers = grid[0] if grid else []
        data_rows = grid[1:] if len(grid) > 1 else []
        serialized = _serialize_table(headers, data_rows, page_idx, t_idx)
        result.append(TableData(page=page_idx, table_index=t_idx,
                                headers=headers, rows=data_rows, serialized=serialized))
    return result


def _extract_pdf(source: str | bytes | Path) -> ExtractionResult:
    if isinstance(source, (str, Path)):
        doc = fitz.open(str(source))
    else:
        doc = fitz.open(stream=source, filetype="pdf")

    body_parts: list[str] = []
    all_tables: list[TableData] = []
    for page_idx, page in enumerate(doc):
        bboxes = _table_bboxes(page)
        body_parts.append(_pdf_body_text(page, bboxes))
        all_tables.extend(_pdf_tables(page, page_idx))
    doc.close()

    body_text = clean_text("\n".join(body_parts))
    table_block = "\n\n".join(t.serialized for t in all_tables)
    combined = (body_text + "\n\n" + table_block).strip() if table_block else body_text
    return ExtractionResult(text=body_text, tables=all_tables, combined=combined)


# ── DOCX extraction ────────────────────────────────────────────────────────

def _docx_tables(doc: Document) -> list[TableData]:
    all_tables: list[TableData] = []
    for t_idx, table in enumerate(doc.tables):
        grid = [[_cell(cell.text) for cell in row.cells] for row in table.rows]
        if not grid:
            continue
        # Deduplicate merged cells (python-docx repeats merged cell text)
        deduped: list[list[str]] = []
        for row in grid:
            deduped_row: list[str] = []
            prev = object()
            for val in row:
                deduped_row.append(val if val != prev else "")
                prev = val
            deduped.append(deduped_row)

        headers = deduped[0] if deduped else []
        data_rows = deduped[1:] if len(deduped) > 1 else []
        serialized = _serialize_table(headers, data_rows, 0, t_idx)
        all_tables.append(TableData(page=0, table_index=t_idx,
                                    headers=headers, rows=data_rows, serialized=serialized))
    return all_tables


def _docx_body_text(doc: Document) -> str:
    """Extract paragraph text, skipping paragraphs that are inside tables."""
    # Collect all paragraph XML elements that live inside a table
    table_para_ids: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_para_ids.add(id(para))

    lines: list[str] = []
    for para in doc.paragraphs:
        if id(para) in table_para_ids:
            continue
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def _extract_docx(source: str | bytes | Path) -> ExtractionResult:
    if isinstance(source, (str, Path)):
        doc = Document(str(source))
    else:
        doc = Document(io.BytesIO(source))

    body_text = clean_text(_docx_body_text(doc))
    all_tables = _docx_tables(doc)
    table_block = "\n\n".join(t.serialized for t in all_tables)
    combined = (body_text + "\n\n" + table_block).strip() if table_block else body_text
    return ExtractionResult(text=body_text, tables=all_tables, combined=combined)


# ── Public API ─────────────────────────────────────────────────────────────

def _is_docx(source: str | bytes | Path, filename: str | None) -> bool:
    """Detect docx by filename extension or by magic bytes (PK zip header)."""
    if filename and filename.lower().endswith(".docx"):
        return True
    if isinstance(source, bytes):
        return source[:2] == b"PK"   # docx is a zip file
    if isinstance(source, Path):
        return source.suffix.lower() == ".docx"
    return False


def extract_full(
    source: str | bytes | Path,
    filename: str | None = None,
) -> ExtractionResult:
    """
    Auto-detect PDF or DOCX and extract body text + tables.
    Returns an ExtractionResult whose `.combined` is ready for embedding.
    """
    if _is_docx(source, filename):
        return _extract_docx(source)
    return _extract_pdf(source)


def extract_and_clean(
    source: str | bytes | Path,
    filename: str | None = None,
) -> str:
    """Convenience wrapper — returns combined text ready for embed_text()."""
    return extract_full(source, filename).combined


def chunk_text(text: str, chunk_size: int = 200, overlap: int = 32) -> list[str]:
    words = text.split()
    if not words:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += chunk_size - overlap
    return [c for c in chunks if c.strip()]
